from flask import Flask, render_template, request, send_file, jsonify
from flask_cors import CORS
from PIL import Image
from io import BytesIO
import os
import logging
import traceback
import re
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import fitz  # PyMuPDF
import datetime
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import A4
from reportlab.lib.utils import simpleSplit
import tempfile
import PyPDF2
import pdfplumber
import pandas as pd

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

app = Flask(__name__)
CORS(app)

# Configure maximum file size (50MB)
app.config['MAX_CONTENT_LENGTH'] = 50 * 1024 * 1024
app.config['SECRET_KEY'] = 'your-secret-key-here'

# Ensure templates folder exists
os.makedirs('templates', exist_ok=True)

@app.route("/")
def home():
    return render_template("index.html")

@app.route("/test", methods=["GET"])
def test():
    return jsonify({"status": "Server is running"}), 200

# ==============================
# 🔹 PDF PAGE ADD (Insert Pages from Another PDF)
# ==============================
@app.route("/pdf-add-pages", methods=["POST", "OPTIONS"])
def pdf_add_pages():
    if request.method == "OPTIONS":
        return _build_cors_preflight_response()
    
    try:
        logger.info("Received pdf-add-pages request")
        
        if 'base_pdf' not in request.files:
            return "No base PDF file provided", 400
            
        if 'insert_pdf' not in request.files:
            return "No insert PDF file provided", 400
            
        base_file = request.files.get("base_pdf")
        insert_file = request.files.get("insert_pdf")
        position_str = request.form.get("position", "")
        
        if not base_file or base_file.filename == "":
            return "No base PDF file selected", 400
            
        if not insert_file or insert_file.filename == "":
            return "No insert PDF file selected", 400
            
        if not base_file.filename.lower().endswith('.pdf'):
            return "Base file must be a PDF", 400
            
        if not insert_file.filename.lower().endswith('.pdf'):
            return "Insert file must be a PDF", 400
            
        if not position_str.strip():
            return "Please specify insertion position", 400
        
        try:
            position = int(position_str)
            if position < 0:
                return "Position must be 0 or greater", 400
        except ValueError:
            return "Invalid position. Please enter a number", 400
        
        logger.info(f"Inserting pages at position {position}")
        
        # Read both PDFs
        base_pdf_bytes = base_file.read()
        insert_pdf_bytes = insert_file.read()
        
        base_stream = BytesIO(base_pdf_bytes)
        insert_stream = BytesIO(insert_pdf_bytes)
        
        # Open PDFs
        base_reader = PyPDF2.PdfReader(base_stream)
        insert_reader = PyPDF2.PdfReader(insert_stream)
        
        total_base_pages = len(base_reader.pages)
        total_insert_pages = len(insert_reader.pages)
        
        logger.info(f"Base PDF has {total_base_pages} pages")
        logger.info(f"Insert PDF has {total_insert_pages} pages")
        
        # Validate position
        if position > total_base_pages:
            return f"Position {position} exceeds total pages ({total_base_pages})", 400
        
        # Create PDF writer
        pdf_writer = PyPDF2.PdfWriter()
        
        # Add pages before insertion point
        for page_num in range(position):
            pdf_writer.add_page(base_reader.pages[page_num])
        
        # Add all pages from insert PDF
        for page_num in range(total_insert_pages):
            pdf_writer.add_page(insert_reader.pages[page_num])
        
        # Add remaining pages from base PDF
        for page_num in range(position, total_base_pages):
            pdf_writer.add_page(base_reader.pages[page_num])
        
        # Create output PDF
        output_pdf = BytesIO()
        pdf_writer.write(output_pdf)
        output_pdf.seek(0)
        
        # Prepare response
        base_name = os.path.splitext(base_file.filename)[0]
        output_filename = f"{base_name}_with_inserted_pages.pdf"
        
        response = send_file(
            output_pdf,
            as_attachment=True,
            download_name=output_filename,
            mimetype="application/pdf"
        )
        response.headers.add('Access-Control-Allow-Origin', '*')
        return response
        
    except Exception as e:
        logger.error(f"Error in pdf_add_pages: {str(e)}")
        logger.error(traceback.format_exc())
        return f"Error adding pages: {str(e)}", 500

# ==============================
# 🔹 PDF TO EXCEL (Enhanced - Better Table Detection)
# ==============================
@app.route("/pdf-to-excel", methods=["POST", "OPTIONS"])
def pdf_to_excel():
    if request.method == "OPTIONS":
        return _build_cors_preflight_response()
    
    try:
        logger.info("Received pdf-to-excel request")
        
        if 'pdf' not in request.files:
            return "No PDF file provided", 400
            
        file = request.files.get("pdf")
        
        if not file or file.filename == "":
            return "No file selected", 400
            
        if not file.filename.lower().endswith('.pdf'):
            return "Please upload a PDF file", 400
            
        logger.info(f"Processing PDF for tables: {file.filename}")
        
        # Read PDF
        pdf_bytes = file.read()
        pdf_stream = BytesIO(pdf_bytes)
        
        # Extract tables using pdfplumber with enhanced settings
        all_tables = []
        table_count = 0
        extracted_text_for_debug = []
        
        with pdfplumber.open(pdf_stream) as pdf:
            for page_num, page in enumerate(pdf.pages, 1):
                logger.info(f"Processing page {page_num}")
                
                # Method 1: Standard table extraction
                tables = page.extract_tables()
                
                if tables and len(tables) > 0:
                    for table_idx, table in enumerate(tables):
                        if table and len(table) > 1:
                            # Clean the table data
                            cleaned_table = []
                            for row in table:
                                cleaned_row = []
                                for cell in row:
                                    if cell is None or cell == "":
                                        cleaned_row.append("")
                                    else:
                                        cell_text = str(cell).strip()
                                        cell_text = re.sub(r'\n', ' ', cell_text)
                                        cell_text = re.sub(r'\s+', ' ', cell_text)
                                        cleaned_row.append(cell_text)
                                
                                if any(cell.strip() for cell in cleaned_row):
                                    cleaned_table.append(cleaned_row)
                            
                            if cleaned_table and len(cleaned_table) > 0:
                                all_tables.append({
                                    'page': page_num,
                                    'table': table_idx + 1,
                                    'data': cleaned_table,
                                    'method': 'standard'
                                })
                                table_count += 1
                                logger.info(f"Found table {table_idx + 1} on page {page_num} (standard method)")
                
                # Method 2: Try to detect tables using lines and text positions
                if not tables or len(tables) == 0:
                    logger.info(f"Trying alternative detection for page {page_num}")
                    
                    words = page.extract_words()
                    
                    if words:
                        rows = {}
                        for word in words:
                            y_pos = round(word['top'], 1)
                            if y_pos not in rows:
                                rows[y_pos] = []
                            rows[y_pos].append(word)
                        
                        sorted_rows = sorted(rows.items())
                        
                        if len(sorted_rows) > 1:
                            table_data = []
                            for y_pos, row_words in sorted_rows:
                                row_words.sort(key=lambda w: w['x0'])
                                row_text = [w['text'] for w in row_words]
                                table_data.append(row_text)
                            
                            if table_data and len(table_data) > 1:
                                all_tables.append({
                                    'page': page_num,
                                    'table': 1,
                                    'data': table_data,
                                    'method': 'text_based'
                                })
                                table_count += 1
                                logger.info(f"Found text-based table on page {page_num}")
                
                # Method 3: Try to extract text in columns
                if not tables or len(tables) == 0:
                    page_text = page.extract_text()
                    if page_text:
                        lines = page_text.split('\n')
                        potential_table = []
                        
                        for line in lines:
                            if '  ' in line and len(line.split()) > 2:
                                columns = re.split(r'\s{2,}', line)
                                if len(columns) > 1:
                                    potential_table.append([col.strip() for col in columns])
                        
                        if len(potential_table) > 1:
                            all_tables.append({
                                'page': page_num,
                                'table': 1,
                                'data': potential_table,
                                'method': 'column_detection'
                            })
                            table_count += 1
                            logger.info(f"Found column-based table on page {page_num}")
                
                # Extract text for debugging
                page_text = page.extract_text()
                if page_text:
                    extracted_text_for_debug.append(f"Page {page_num} text (first 500 chars): {page_text[:500]}")
        
        if table_count == 0:
            error_msg = "No tables found in the PDF. "
            error_msg += "This could be because:\n"
            error_msg += "• The PDF contains scanned images (needs OCR)\n"
            error_msg += "• Tables are in image format\n"
            error_msg += "• Tables don't have borders/lines\n"
            error_msg += "• The PDF uses complex formatting\n\n"
            
            if extracted_text_for_debug:
                error_msg += "Debug info (first few pages):\n"
                error_msg += "\n".join(extracted_text_for_debug[:3])
            
            return error_msg, 400
        
        logger.info(f"Found {table_count} tables in total")
        
        # Create Excel file with multiple sheets
        output_excel = BytesIO()
        
        with pd.ExcelWriter(output_excel, engine='openpyxl') as writer:
            for i, table_info in enumerate(all_tables):
                data = table_info['data']
                
                if len(data) > 0:
                    first_row = data[0]
                    has_headers = False
                    
                    if first_row and all(str(cell).strip() for cell in first_row):
                        has_numbers = any(str(cell).replace('.', '').replace('-', '').isdigit() for cell in first_row)
                        if not has_numbers:
                            has_headers = True
                    
                    if has_headers and len(data) > 1:
                        df = pd.DataFrame(data[1:], columns=[str(col).strip() if col else f"Column_{j+1}" for j, col in enumerate(first_row)])
                    else:
                        df = pd.DataFrame(data)
                        df.columns = [f"Column_{j+1}" for j in range(len(df.columns))]
                    
                    sheet_name = f"Page_{table_info['page']}_Table_{table_info['table']}"
                    sheet_name = sheet_name[:31]
                    
                    df.to_excel(writer, sheet_name=sheet_name, index=False)
                    
                    # Auto-adjust column widths
                    worksheet = writer.sheets[sheet_name]
                    for column in df:
                        column_width = max(df[column].astype(str).map(len).max(), len(column))
                        column_width = min(column_width, 50)
                        col_idx = df.columns.get_loc(column)
                        worksheet.column_dimensions[chr(65 + col_idx)].width = column_width + 2
        
        output_excel.seek(0)
        
        original_name = os.path.splitext(file.filename)[0]
        output_filename = f"{original_name}_tables.xlsx"
        
        response = send_file(
            output_excel,
            as_attachment=True,
            download_name=output_filename,
            mimetype="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
        )
        response.headers.add('Access-Control-Allow-Origin', '*')
        return response
        
    except Exception as e:
        logger.error(f"Error in pdf_to_excel: {str(e)}")
        logger.error(traceback.format_exc())
        return f"Error converting PDF to Excel: {str(e)}", 500

# ==============================
# 🔹 PDF PAGE DELETE
# ==============================
@app.route("/pdf-delete-pages", methods=["POST", "OPTIONS"])
def pdf_delete_pages():
    if request.method == "OPTIONS":
        return _build_cors_preflight_response()
    
    try:
        logger.info("Received pdf-delete-pages request")
        
        if 'pdf' not in request.files:
            return "No PDF file provided", 400
            
        file = request.files.get("pdf")
        
        if not file or file.filename == "":
            return "No file selected", 400
            
        if not file.filename.lower().endswith('.pdf'):
            return "Please upload a PDF file", 400
            
        pages_to_delete_str = request.form.get("pages", "")
        
        if not pages_to_delete_str.strip():
            return "Please provide page numbers to delete", 400
        
        try:
            pages_to_delete = set()
            for part in pages_to_delete_str.split(','):
                part = part.strip()
                if '-' in part:
                    start, end = map(int, part.split('-'))
                    pages_to_delete.update(range(start, end + 1))
                else:
                    pages_to_delete.add(int(part))
        except ValueError:
            return "Invalid page numbers format. Use comma-separated numbers (e.g., 1,3,5) or ranges (e.g., 1-5)", 400
        
        if not pages_to_delete:
            return "No valid page numbers provided", 400
        
        pdf_bytes = file.read()
        pdf_stream = BytesIO(pdf_bytes)
        pdf_reader = PyPDF2.PdfReader(pdf_stream)
        total_pages = len(pdf_reader.pages)
        
        invalid_pages = [p for p in pages_to_delete if p < 1 or p > total_pages]
        if invalid_pages:
            return f"Invalid page numbers: {invalid_pages}. PDF has {total_pages} pages.", 400
        
        pdf_writer = PyPDF2.PdfWriter()
        pages_added = 0
        
        for page_num in range(1, total_pages + 1):
            if page_num not in pages_to_delete:
                pdf_writer.add_page(pdf_reader.pages[page_num - 1])
                pages_added += 1
        
        if pages_added == 0:
            return "Cannot delete all pages. At least one page must remain.", 400
        
        output_pdf = BytesIO()
        pdf_writer.write(output_pdf)
        output_pdf.seek(0)
        
        original_name = os.path.splitext(file.filename)[0]
        output_filename = f"{original_name}_deleted_pages.pdf"
        
        response = send_file(
            output_pdf,
            as_attachment=True,
            download_name=output_filename,
            mimetype="application/pdf"
        )
        response.headers.add('Access-Control-Allow-Origin', '*')
        return response
        
    except Exception as e:
        logger.error(f"Error: {str(e)}")
        return f"Error deleting pages: {str(e)}", 500

# ==============================
# 🔹 PDF PAGE COPY
# ==============================
@app.route("/pdf-copy-pages", methods=["POST", "OPTIONS"])
def pdf_copy_pages():
    if request.method == "OPTIONS":
        return _build_cors_preflight_response()
    
    try:
        logger.info("Received pdf-copy-pages request")
        
        if 'pdf' not in request.files:
            return "No PDF file provided", 400
            
        file = request.files.get("pdf")
        
        if not file or file.filename == "":
            return "No file selected", 400
            
        if not file.filename.lower().endswith('.pdf'):
            return "Please upload a PDF file", 400
            
        pages_to_copy_str = request.form.get("pages", "")
        
        if not pages_to_copy_str.strip():
            return "Please provide page numbers to extract", 400
        
        try:
            pages_to_copy = []
            for part in pages_to_copy_str.split(','):
                part = part.strip()
                if '-' in part:
                    start, end = map(int, part.split('-'))
                    pages_to_copy.extend(range(start, end + 1))
                else:
                    pages_to_copy.append(int(part))
        except ValueError:
            return "Invalid page numbers format. Use comma-separated numbers (e.g., 1,3,5) or ranges (e.g., 1-5)", 400
        
        if not pages_to_copy:
            return "No valid page numbers provided", 400
        
        pages_to_copy = list(dict.fromkeys(pages_to_copy))
        
        pdf_bytes = file.read()
        pdf_stream = BytesIO(pdf_bytes)
        pdf_reader = PyPDF2.PdfReader(pdf_stream)
        total_pages = len(pdf_reader.pages)
        
        invalid_pages = [p for p in pages_to_copy if p < 1 or p > total_pages]
        if invalid_pages:
            return f"Invalid page numbers: {invalid_pages}. PDF has {total_pages} pages.", 400
        
        pdf_writer = PyPDF2.PdfWriter()
        
        for page_num in pages_to_copy:
            pdf_writer.add_page(pdf_reader.pages[page_num - 1])
        
        output_pdf = BytesIO()
        pdf_writer.write(output_pdf)
        output_pdf.seek(0)
        
        original_name = os.path.splitext(file.filename)[0]
        output_filename = f"{original_name}_extracted_pages.pdf"
        
        response = send_file(
            output_pdf,
            as_attachment=True,
            download_name=output_filename,
            mimetype="application/pdf"
        )
        response.headers.add('Access-Control-Allow-Origin', '*')
        return response
        
    except Exception as e:
        logger.error(f"Error: {str(e)}")
        return f"Error copying pages: {str(e)}", 500

# ==============================
# 🔹 PDF MERGE
# ==============================
@app.route("/pdf-merge", methods=["POST", "OPTIONS"])
def pdf_merge():
    if request.method == "OPTIONS":
        return _build_cors_preflight_response()
    
    try:
        logger.info("Received pdf-merge request")
        
        if 'pdfs' not in request.files:
            return "No PDF files provided", 400
            
        files = request.files.getlist("pdfs")
        
        if not files or len(files) == 0:
            return "No files selected", 400
        
        valid_files = []
        invalid_files = []
        
        for file in files:
            if file and file.filename:
                if file.filename.lower().endswith('.pdf'):
                    valid_files.append(file)
                else:
                    invalid_files.append(file.filename)
        
        if not valid_files:
            return "No valid PDF files found. Please upload PDF files.", 400
        
        if invalid_files:
            logger.warning(f"Skipped non-PDF files: {invalid_files}")
        
        logger.info(f"Merging {len(valid_files)} PDF files")
        
        pdf_merger = PyPDF2.PdfMerger()
        failed_files = []
        successful_count = 0
        
        for file in valid_files:
            try:
                pdf_bytes = file.read()
                pdf_stream = BytesIO(pdf_bytes)
                pdf_merger.append(pdf_stream)
                successful_count += 1
                logger.info(f"Successfully added: {file.filename}")
                
            except PyPDF2.errors.PdfReadError as e:
                logger.error(f"Corrupted PDF {file.filename}: {str(e)}")
                failed_files.append(file.filename)
            except Exception as e:
                logger.error(f"Error processing {file.filename}: {str(e)}")
                failed_files.append(file.filename)
        
        if successful_count == 0:
            return "No valid PDF files could be processed. All files are corrupted or invalid.", 400
        
        output_pdf = BytesIO()
        pdf_merger.write(output_pdf)
        pdf_merger.close()
        output_pdf.seek(0)
        
        timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
        output_filename = f"merged_{timestamp}.pdf"
        
        if failed_files:
            logger.warning(f"Failed to merge {len(failed_files)} files: {failed_files}")
        
        response = send_file(
            output_pdf,
            as_attachment=True,
            download_name=output_filename,
            mimetype="application/pdf"
        )
        response.headers.add('Access-Control-Allow-Origin', '*')
        return response
        
    except Exception as e:
        logger.error(f"Error: {str(e)}")
        return f"Error merging PDFs: {str(e)}", 500

# ==============================
# 🔹 WORD → PDF
# ==============================
@app.route("/word-to-pdf", methods=["POST", "OPTIONS"])
def word_to_pdf():
    if request.method == "OPTIONS":
        return _build_cors_preflight_response()
    
    try:
        logger.info("Received word-to-pdf request")
        
        if 'word' not in request.files:
            return "No Word file provided", 400
            
        file = request.files.get("word")
        
        if not file or file.filename == "":
            return "No file selected", 400
            
        if not file.filename.lower().endswith('.docx'):
            return "Please upload a Word (.docx) file", 400
            
        logger.info(f"Processing Word document: {file.filename}")
        
        doc = Document(file)
        
        temp_pdf = tempfile.NamedTemporaryFile(delete=False, suffix='.pdf')
        temp_pdf.close()
        
        c = canvas.Canvas(temp_pdf.name, pagesize=A4)
        width, height = A4
        
        left_margin = 72
        right_margin = width - 72
        top_margin = height - 72
        bottom_margin = 72
        y_position = top_margin
        line_height = 14
        default_font_size = 11
        
        def add_new_page():
            nonlocal y_position
            c.showPage()
            y_position = top_margin
            c.setFont("Helvetica", default_font_size)
        
        def check_page_break(needed_space):
            nonlocal y_position
            if y_position - needed_space < bottom_margin:
                add_new_page()
        
        def add_paragraph(paragraph):
            nonlocal y_position
            
            text = paragraph.text.strip()
            if not text:
                check_page_break(line_height / 2)
                y_position -= line_height / 2
                return
            
            font_size = default_font_size
            bold = False
            italic = False
            alignment = 'left'
            
            if paragraph.style:
                style_name = paragraph.style.name.lower()
                
                if 'heading 1' in style_name or 'heading1' in style_name:
                    font_size = 18
                    bold = True
                    alignment = 'center'
                elif 'heading 2' in style_name or 'heading2' in style_name:
                    font_size = 16
                    bold = True
                elif 'heading 3' in style_name or 'heading3' in style_name:
                    font_size = 14
                    bold = True
                elif 'heading' in style_name:
                    font_size = 12
                    bold = True
                
                if paragraph.alignment and paragraph.alignment == 1:
                    alignment = 'center'
                elif paragraph.alignment and paragraph.alignment == 2:
                    alignment = 'right'
            
            if text.isupper() and len(text) < 100 and not bold:
                font_size = 12
                bold = True
            
            runs_info = []
            for run in paragraph.runs:
                if run.text:
                    runs_info.append({
                        'text': run.text,
                        'bold': run.bold,
                        'italic': run.italic,
                        'underline': run.underline
                    })
            
            if not runs_info:
                return
            
            check_page_break(line_height)
            y_position -= line_height / 2
            
            if len(runs_info) > 1:
                x_position = left_margin
                for run_info in runs_info:
                    run_text = run_info['text']
                    if not run_text:
                        continue
                    
                    font_name = "Helvetica"
                    if run_info['bold'] and run_info['italic']:
                        font_name = "Helvetica-BoldOblique"
                    elif run_info['bold']:
                        font_name = "Helvetica-Bold"
                    elif run_info['italic']:
                        font_name = "Helvetica-Oblique"
                    
                    c.setFont(font_name, font_size)
                    text_width = c.stringWidth(run_text, font_name, font_size)
                    
                    if y_position < bottom_margin:
                        add_new_page()
                        x_position = left_margin
                    
                    if run_info['underline']:
                        c.drawString(x_position, y_position, run_text)
                        c.line(x_position, y_position - 2, x_position + text_width, y_position - 2)
                    else:
                        c.drawString(x_position, y_position, run_text)
                    
                    x_position += text_width
                
                y_position -= line_height
            else:
                max_width = right_margin - left_margin
                font_name = "Helvetica"
                if bold and italic:
                    font_name = "Helvetica-BoldOblique"
                elif bold:
                    font_name = "Helvetica-Bold"
                elif italic:
                    font_name = "Helvetica-Oblique"
                
                c.setFont(font_name, font_size)
                
                words = text.split(' ')
                lines = []
                current_line = ""
                
                for word in words:
                    test_line = current_line + " " + word if current_line else word
                    if c.stringWidth(test_line, font_name, font_size) <= max_width:
                        current_line = test_line
                    else:
                        if current_line:
                            lines.append(current_line)
                        current_line = word
                
                if current_line:
                    lines.append(current_line)
                
                for i, line in enumerate(lines):
                    if i > 0:
                        check_page_break(line_height)
                    
                    line_width = c.stringWidth(line, font_name, font_size)
                    
                    if alignment == 'center':
                        x = left_margin + (max_width - line_width) / 2
                    elif alignment == 'right':
                        x = right_margin - line_width
                    else:
                        x = left_margin
                    
                    has_underline = len(runs_info) > 0 and runs_info[0].get('underline', False)
                    
                    if has_underline:
                        c.drawString(x, y_position, line)
                        c.line(x, y_position - 2, x + line_width, y_position - 2)
                    else:
                        c.drawString(x, y_position, line)
                    
                    y_position -= line_height
        
        def add_table(table):
            nonlocal y_position
            
            if not table.rows:
                return
            
            check_page_break(line_height * 3)
            y_position -= line_height
            
            num_cols = len(table.rows[0].cells)
            col_widths = []
            
            for col in range(num_cols):
                max_width = 0
                for row in table.rows:
                    cell_text = row.cells[col].text.strip()
                    if cell_text:
                        text_width = c.stringWidth(cell_text, "Helvetica", 10)
                        max_width = max(max_width, text_width)
                col_widths.append(min(max_width + 20, (right_margin - left_margin) / num_cols))
            
            row_height = 20
            
            for col, width in enumerate(col_widths):
                x = left_margin + sum(col_widths[:col])
                c.rect(x, y_position - row_height, width, row_height, stroke=1, fill=0)
                
                if len(table.rows) > 0:
                    header_text = table.rows[0].cells[col].text.strip()
                    if header_text:
                        c.setFont("Helvetica-Bold", 10)
                        text_width = c.stringWidth(header_text, "Helvetica-Bold", 10)
                        text_x = x + (width - text_width) / 2
                        c.drawString(text_x, y_position - row_height + 5, header_text)
            
            y_position -= row_height
            
            for row_idx, row in enumerate(table.rows[1:], 1):
                if y_position - row_height < bottom_margin:
                    add_new_page()
                    for col, width in enumerate(col_widths):
                        x = left_margin + sum(col_widths[:col])
                        c.rect(x, y_position - row_height, width, row_height, stroke=1, fill=0)
                        if len(table.rows) > 0:
                            header_text = table.rows[0].cells[col].text.strip()
                            if header_text:
                                c.setFont("Helvetica-Bold", 10)
                                text_width = c.stringWidth(header_text, "Helvetica-Bold", 10)
                                text_x = x + (width - text_width) / 2
                                c.drawString(text_x, y_position - row_height + 5, header_text)
                    y_position -= row_height
                
                for col, width in enumerate(col_widths):
                    x = left_margin + sum(col_widths[:col])
                    c.rect(x, y_position - row_height, width, row_height, stroke=1, fill=0)
                    
                    cell_text = row.cells[col].text.strip()
                    if cell_text:
                        c.setFont("Helvetica", 9)
                        lines = simpleSplit(cell_text, "Helvetica", 9, width - 10)
                        for i, line in enumerate(lines[:2]):
                            if i < 2:
                                c.drawString(x + 5, y_position - row_height + 5 + (i * 10), line)
                
                y_position -= row_height
            
            y_position -= line_height
        
        for element in doc.element.body:
            if element.tag.endswith('p'):
                for paragraph in doc.paragraphs:
                    if paragraph._element is element:
                        add_paragraph(paragraph)
                        break
            elif element.tag.endswith('tbl'):
                for table in doc.tables:
                    if table._element is element:
                        add_table(table)
                        break
        
        c.save()
        
        with open(temp_pdf.name, 'rb') as f:
            pdf_bytes = f.read()
        
        os.unlink(temp_pdf.name)
        
        pdf_io = BytesIO(pdf_bytes)
        pdf_io.seek(0)
        
        original_name = os.path.splitext(file.filename)[0]
        
        response = send_file(
            pdf_io,
            as_attachment=True,
            download_name=f"{original_name}.pdf",
            mimetype="application/pdf"
        )
        response.headers.add('Access-Control-Allow-Origin', '*')
        return response
        
    except Exception as e:
        logger.error(f"Error: {str(e)}")
        return f"Error converting Word to PDF: {str(e)}", 500

# ==============================
# 🔹 IMAGE → PDF
# ==============================
@app.route("/image-to-pdf", methods=["POST", "OPTIONS"])
def image_to_pdf():
    if request.method == "OPTIONS":
        return _build_cors_preflight_response()
    
    try:
        if 'image' not in request.files:
            return "No image file provided", 400
            
        file = request.files.get("image")
        
        if not file or file.filename == "":
            return "No file selected", 400

        img = Image.open(file)
        
        if img.mode in ("RGBA", "LA", "P"):
            if img.mode == "RGBA":
                background = Image.new("RGB", img.size, (255, 255, 255))
                background.paste(img, mask=img.split()[3])
                img = background
            else:
                img = img.convert("RGB")
        elif img.mode != "RGB":
            img = img.convert("RGB")

        pdf_io = BytesIO()
        img.save(pdf_io, format="PDF", resolution=150.0, optimize=True)
        pdf_io.seek(0)

        response = send_file(
            pdf_io,
            as_attachment=True,
            download_name="converted.pdf",
            mimetype="application/pdf"
        )
        response.headers.add('Access-Control-Allow-Origin', '*')
        return response

    except Exception as e:
        logger.error(f"Error: {str(e)}")
        return f"Error converting image to PDF: {str(e)}", 500

# ==============================
# 🔹 MULTIPLE IMAGES → PDF
# ==============================
@app.route("/multi-image-to-pdf", methods=["POST", "OPTIONS"])
def multi_image_to_pdf():
    if request.method == "OPTIONS":
        return _build_cors_preflight_response()
    
    try:
        files = request.files.getlist("images")

        if not files or len(files) == 0:
            return "No files selected", 400

        files = [f for f in files if f and f.filename]
        
        if not files:
            return "No valid files selected", 400

        image_list = []

        for file in files:
            try:
                img = Image.open(file)
                
                if img.mode in ("RGBA", "LA", "P"):
                    if img.mode == "RGBA":
                        background = Image.new("RGB", img.size, (255, 255, 255))
                        background.paste(img, mask=img.split()[3])
                        img = background
                    else:
                        img = img.convert("RGB")
                elif img.mode != "RGB":
                    img = img.convert("RGB")
                
                image_list.append(img)
                
            except Exception as e:
                logger.error(f"Failed to load image {file.filename}: {str(e)}")

        if not image_list:
            return "No valid images could be processed", 400

        pdf_io = BytesIO()
        
        first_image = image_list[0]
        if len(image_list) > 1:
            first_image.save(
                pdf_io, 
                format="PDF", 
                resolution=150.0,
                save_all=True, 
                append_images=image_list[1:],
                optimize=True
            )
        else:
            first_image.save(pdf_io, format="PDF", resolution=150.0, optimize=True)
        
        pdf_io.seek(0)

        response = send_file(
            pdf_io,
            as_attachment=True,
            download_name="combined.pdf",
            mimetype="application/pdf"
        )
        response.headers.add('Access-Control-Allow-Origin', '*')
        return response

    except Exception as e:
        logger.error(f"Error: {str(e)}")
        return f"Error converting images to PDF: {str(e)}", 500

# ==============================
# 🔹 CROP IMAGE
# ==============================
@app.route("/crop-image", methods=["POST", "OPTIONS"])
def crop_image():
    if request.method == "OPTIONS":
        return _build_cors_preflight_response()
    
    try:
        file = request.files.get("image")

        if not file or file.filename == "":
            return "No file selected", 400

        try:
            x = int(request.form.get("x", 0))
            y = int(request.form.get("y", 0))
            w = int(request.form.get("width", 0))
            h = int(request.form.get("height", 0))
        except ValueError:
            return "Invalid crop coordinates", 400

        if w <= 0 or h <= 0:
            return "Invalid width or height for cropping", 400

        img = Image.open(file)
        
        if x < 0 or y < 0:
            return "Crop coordinates cannot be negative.", 400
            
        if x + w > img.width or y + h > img.height:
            return f"Crop area exceeds image size ({img.width}x{img.height}).", 400

        cropped = img.crop((x, y, x + w, y + h))
        
        if cropped.mode in ("RGBA", "LA", "P"):
            if cropped.mode == "RGBA":
                background = Image.new("RGB", cropped.size, (255, 255, 255))
                background.paste(cropped, mask=cropped.split()[3])
                cropped = background
            else:
                cropped = cropped.convert("RGB")
        elif cropped.mode != "RGB":
            cropped = cropped.convert("RGB")

        img_bytes = BytesIO()
        cropped.save(img_bytes, format="JPEG", quality=95, optimize=True)
        img_bytes.seek(0)

        response = send_file(
            img_bytes,
            as_attachment=True,
            download_name="cropped.jpg",
            mimetype="image/jpeg"
        )
        response.headers.add('Access-Control-Allow-Origin', '*')
        return response

    except Exception as e:
        logger.error(f"Error: {str(e)}")
        return f"Error cropping image: {str(e)}", 500

# ==============================
# 🔹 PDF → WORD
# ==============================
@app.route("/pdf-to-word", methods=["POST", "OPTIONS"])
def pdf_to_word():
    if request.method == "OPTIONS":
        return _build_cors_preflight_response()
    
    try:
        if 'pdf' not in request.files:
            return "No PDF file provided", 400
            
        file = request.files.get("pdf")
        
        if not file or file.filename == "":
            return "No file selected", 400
            
        if not file.filename.lower().endswith('.pdf'):
            return "Please upload a PDF file", 400
            
        pdf_bytes = file.read()
        pdf_document = fitz.open(stream=pdf_bytes, filetype="pdf")
        
        doc = Document()
        
        for page_num in range(len(pdf_document)):
            page = pdf_document.load_page(page_num)
            text_lines = page.get_text().split('\n')
            
            for line in text_lines:
                if line.strip():
                    paragraph = doc.add_paragraph()
                    run = paragraph.add_run(line)
                    run.font.size = Pt(11)
            
            if page_num < len(pdf_document) - 1:
                doc.add_page_break()
        
        pdf_document.close()
        
        doc_bytes = BytesIO()
        doc.save(doc_bytes)
        doc_bytes.seek(0)
        
        original_name = os.path.splitext(file.filename)[0]
        
        response = send_file(
            doc_bytes,
            as_attachment=True,
            download_name=f"{original_name}.docx",
            mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
        response.headers.add('Access-Control-Allow-Origin', '*')
        return response
        
    except Exception as e:
        logger.error(f"Error: {str(e)}")
        return f"Error converting PDF to Word: {str(e)}", 500

def _build_cors_preflight_response():
    response = jsonify({"status": "preflight"})
    response.headers.add("Access-Control-Allow-Origin", "*")
    response.headers.add('Access-Control-Allow-Headers', "*")
    response.headers.add('Access-Control-Allow-Methods', "*")
    return response, 200

@app.errorhandler(413)
def too_large(e):
    return "File too large. Maximum size is 50MB.", 413

if __name__ == "__main__":
    print("\n" + "="*50)
    print("📄 CONVERTER SUITE - FLASK SERVER")
    print("="*50)
    print("\n=== Available Features ===")
    print("  ✅ Image → PDF")
    print("  ✅ Multiple Images → PDF")
    print("  ✅ Crop Image")
    print("  ✅ PDF → Word")
    print("  ✅ Word → PDF")
    print("  ✅ PDF Merge")
    print("  ✅ PDF Page Delete")
    print("  ✅ PDF Page Copy")
    print("  ✅ PDF Page Add")
    print("  ✅ PDF to Excel (Enhanced Table Detection)")
    print("\n🚀 Starting Flask server at http://127.0.0.1:5000")
    print("   Press CTRL+C to stop\n")
    
    app.run(debug=True, host='0.0.0.0', port=5000)